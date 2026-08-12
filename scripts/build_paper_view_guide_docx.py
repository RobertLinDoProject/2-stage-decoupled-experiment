from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "Paper_View_指標白話說明.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "5D6875"
RED = "9B1C1C"
GOLD = "7A5A00"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C7D0DB", size="6") -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_widths(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell_properties = cell._tc.get_or_add_tcPr()
            width_node = cell_properties.first_child_found_in("w:tcW")
            if width_node is None:
                width_node = OxmlElement("w:tcW")
                cell_properties.append(width_node)
            width_node.set(qn("w:w"), str(widths_dxa[index]))
            width_node.set(qn("w:type"), "dxa")


def set_run_font(run, *, size=10.5, color="1F2933", bold=False, italic=False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, *, before=0, after=6, line=1.25, alignment=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(paragraph, text: str, **kwargs) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)


def add_body(doc, text: str, *, after=6, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(paragraph, bold_prefix, bold=True)
        add_text(paragraph, text[len(bold_prefix):])
    else:
        add_text(paragraph, text)


def add_bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    format_paragraph(paragraph, after=4)
    add_text(paragraph, text)


def add_number(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    format_paragraph(paragraph, after=4)
    add_text(paragraph, text)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    format_paragraph(paragraph, before={1: 18, 2: 14, 3: 10}[level], after={1: 10, 2: 7, 3: 5}[level])
    add_text(paragraph, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_callout(doc, title: str, body: str, *, fill=CALLOUT, title_color=NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    set_table_borders(table, color="D5DDE6", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, after=3)
    add_text(paragraph, title, size=11, color=title_color, bold=True)
    paragraph = cell.add_paragraph()
    format_paragraph(paragraph, after=0)
    add_text(paragraph, body, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    set_table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        format_paragraph(paragraph, after=0, line=1.15)
        add_text(paragraph, header, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            format_paragraph(paragraph, after=0, line=1.15)
            add_text(paragraph, value, size=9.2)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format_paragraph(paragraph, after=0, line=1.0)
    add_text(paragraph, "Paper View 指標白話說明 | Decoupled 2-Stage Experiment", size=8.5, color=MUTED)


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    title = document.add_paragraph()
    format_paragraph(title, after=4, line=1.0)
    add_text(title, "Paper View 指標白話說明", size=24, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    format_paragraph(subtitle, after=14, line=1.1)
    add_text(subtitle, "Decoupled 2-Stage Experiment | 讓每個數字都能被看懂", size=12, color=MUTED)

    add_callout(
        document,
        "先記住一件事",
        "Paper View 不是重新計算結果的地方。它只把 M8 已發布的結果依條件篩選並呈現；上方篩選器決定看哪些資料列，欄位 Checkbox 決定每列顯示哪些欄位。",
    )

    add_heading(document, "1. 先看懂一列結果", 1)
    add_body(document, "Paper View 的一列，是某個 rule source、決策方式、framework、密度區間下的一組 M8 aggregate result。先看條件，再看 Branch Valid Rate 與 Executed Trials，最後用 Risk、Action、Failure 指標解釋原因。")
    add_table(document, ["欄位", "白話意思"], [
        ["Rule Source", "M6 使用哪一套拓樸與規則：人工規則或 AI 生成規則。"],
        ["Framework", "w/o 是理想輸入；w/ 是加入 perception residual 後的部署輸入。"],
        ["Trial Type", "ideal 對應 w/o；deployment 對應 w/。"],
        ["決策方式", "M6 由 rule-based planner 做決策，或由 GAI 產生 action。"],
        ["Regime", "M4 scenario 的 ground-truth 人流密度區間：LOW、MEDIUM、HIGH。"],
    ], [1900, 7460])

    add_heading(document, "2. 最重要的三個可靠度數字", 1)
    add_callout(document, "閱讀順序", "先看 R_ideal：正確人數輸入時，決策本身能不能通過 M7。再看 R_deploy：加入 perception 誤差後還剩多少可靠度。最後看 Delta R：兩者差多少。", fill="F4F6F9")
    add_table(document, ["指標", "公式", "白話意思"], [
        ["R_ideal", "ideal branch 的 M7 valid ÷ ideal executed trials", "假設人數辨識完全正確時，決策通過容量、拓樸與規則驗證的比例。"],
        ["R_deploy", "deployment branch 的 M7 valid ÷ deployment executed trials", "人數輸入含 perception residual 時，決策仍通過 M7 的比例。"],
        ["Delta R", "R_ideal - R_deploy", "感知誤差進入部署流程後，可靠度相對理想情況下降多少。"],
        ["Branch Valid Rate", "該列 branch 的 valid rate", "w/o 列等於 R_ideal；w/ 列等於 R_deploy。"],
        ["Executed Trials", "實際有 terminal outcome 的 trials", "這一列的分母是多少。沒有執行的 GAI 不會被假裝成 0。"],
    ], [1500, 2850, 5010])
    add_body(document, "例子：30 個 deployment trials 中有 24 個通過 M7，Branch Valid Rate = 24 ÷ 30 = 0.800。這不是 perception accuracy，而是決策結果通過獨立 M7 驗證的比例。")
    add_callout(document, "R_ideal = 0 的判讀", "代表所有 ideal trials 都沒有通過 M7；這不是 perception residual 造成的。若 R_ideal 與 R_deploy 都是 0，Delta R 數學上雖然是 0，但沒有可解讀的 perception degradation headroom，必須回到 M7 violation evidence。", fill="FFF8E8", title_color=GOLD)

    add_heading(document, "3. Risk 指標：有沒有找對高風險來源", 1)
    add_body(document, "Risk 指標是在看 M6 是否正確辨識哪些 source 需要處理。它不是整體疏散是否成功；整體是否通過仍以 M7 valid 與 Branch Valid Rate 為主。")
    add_table(document, ["指標", "白話意思"], [
        ["Risk Precision", "M6 判定為高風險的來源中，有多少確實是高風險。高表示少誤報。"],
        ["Risk Recall", "實際高風險的來源中，有多少被 M6 找出來。高表示少漏報。"],
        ["Risk Consistency", "把 Precision 與 Recall 合併成一個分數；目前依 F-beta policy 計算。"],
        ["Risk beta", "F-beta 中的權重設定。beta 越大，越重視不要漏掉高風險來源。"],
    ], [2200, 7160])

    add_heading(document, "4. Action 指標：產生的疏散動作是否合規", 1)
    add_table(document, ["指標", "白話意思"], [
        ["Legality", "action 是否使用合法 source、target、edge、容量與人數範圍。"],
        ["Priority", "是否依 M6 規定的 source priority 與 target cost 順序做分配。"],
        ["Economy", "是否避免不必要地把同一來源拆到過多目標，保留較簡潔的分配。"],
        ["Action Consistency", "把 Legality、Priority、Economy 合併看 action 品質的分數；致命 legality failure 會使其失去意義。"],
    ], [2200, 7160])
    add_body(document, "Action 指標可以解釋「為什麼這個 branch 通過或失敗」，但不能取代 M7。M7 仍是最後的獨立驗證標準。")

    add_heading(document, "5. Failure 指標：失敗發生在哪裡", 1)
    add_table(document, ["指標", "白話意思"], [
        ["Invalid Output", "M6 輸出不符合 canonical action 格式或 contract，例如非法欄位、target 或 count。"],
        ["Rule Violation", "M7 發現 action 違反正式規則的 trial 比例。"],
        ["Capacity Violation", "M7 發現某個節點的 post-population 超過容量。"],
        ["Topology Violation", "M7 發現 source → target 不符合正式 topology、allowed edge 或 destination rule。"],
        ["M6 Outcome", "M6 的執行結果：available、invalid_output、decision_infeasible 或 unavailable。"],
        ["Availability", "這組結果是否有可用的正式 terminal outcome。unavailable 不等於 0。"],
        ["Metric Policy", "本次 M8 使用的 metric policy 版本，方便確認不同 Run 是否使用相同計算規則。"],
    ], [2200, 7160])
    add_callout(document, "兩種 0 不一樣", "invalid_output 或 decision_infeasible 是 GAI 已執行但能力／輸出失敗，會以 valid=0 納入分母；unavailable 是沒有正式執行結果，指標應維持 unavailable 或 null，不可轉成 0。", fill="FFF8E8", title_color=GOLD)

    add_heading(document, "6. 三種比較怎麼看", 1)
    add_heading(document, "w/o 與 w/：看 perception 誤差造成的差異", 2)
    add_body(document, "固定 topology、model、regime、rule source 與決策方式後，比較 R_ideal、R_deploy、Delta R，以及 capacity／topology violation。w/o 使用 scenario_gt；w/ 使用加入 residual 後的 observed_population；兩側 M7 都使用同一份 scenario_gt 驗證。")
    add_heading(document, "Rule-based 與 GAI：看決策介面差異", 2)
    add_body(document, "固定 rule source、topology、model、regime 與 framework 後，比較兩種 M6 產生 action 的方式。GAI invalid_output、decision_infeasible 反映模型輸出或決策能力；provider 未執行則是 unavailable，不是失敗分數。")
    add_heading(document, "人工規則與 AI 生成規則：看規則來源差異", 2)
    add_body(document, "固定決策方式與 framework 後，比較 M6 使用人工 rule bundle 或 AI-generated rule bundle。兩者共用相同場景、observation 與人工 M7 gold-standard；因此 AI 規則若允許了不被人工 topology 接受的行動，會在 M7 evidence 中被挑出來。")
    add_heading(document, "7. 最後的閱讀檢查表", 1)
    for item in [
        "先確認目前篩選的 Rule Source、Framework、Trial Type、決策方式與 Regime。",
        "先看 Executed Trials，再看 Branch Valid Rate，避免只看小樣本比例。",
        "用 R_ideal、R_deploy、Delta R 判斷理想與部署差異。",
        "用 Risk 與 Action 指標找原因，用 Failure 指標找違規類型。",
        "看到 unavailable 時，不要當成 0；看到 R_ideal=0 時，回到 M7 evidence。",
        "ALL 只代表顯示全部原始 rows，不代表平均、加總或重新計算。",
    ]:
        add_bullet(document, item)
    add_callout(document, "一句話總結", "Paper View 要回答的是：在相同場景與相同 M7 標準下，哪一種規則來源、決策方式或 framework，讓疏散 action 更容易通過驗證；各個細項指標則用來說明差異發生的原因。")

    add_heading(document, "附註：數值來源", 1)
    add_body(document, "Paper View 的結果直接來自 M8 canonical aggregate rows。M7 提供 trial-level validation evidence；M6 trace 用來說明 GAI 的輸出與執行狀態。這份說明文件不重新計算任何實驗結果，也不把統計結果解讀成未經驗證的因果結論。", after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
